import re
from pathlib import Path
from typing import Dict, List


DEFAULT_JD = {
    "required_skills": [],
    "preferred_skills": [],
    "target_title": "Any",
    "min_experience_years": 0,
    "target_industry": "Any",
    "target_field": "Computer Science",
    "skills_text": "",
}

KNOWN_SKILLS = [
    "Python",
    "SQL",
    "Spark",
    "PySpark",
    "Airflow",
    "Apache Beam",
    "Kafka",
    "AWS",
    "GCP",
    "Azure",
    "Snowflake",
    "BigQuery",
    "Docker",
    "Kubernetes",
    "MLflow",
    "NLP",
    "TensorFlow",
    "PyTorch",
    "Scikit-learn",
    "LLM",
    "Fine-tuning LLMs",
    "React",
    "Next.js",
    "Node.js",
    "Java",
    "TypeScript",
]

TITLE_PATTERNS = [
    r"(?:looking for|hiring|role[:\s]+|position[:\s]+|job title[:\s]+)(?:an?\s+)?([A-Z][A-Za-z /+-]*(?:Engineer|Developer|Scientist|Analyst|Manager|Architect|Specialist))",
    r"\b(ML Engineer|Machine Learning Engineer|Data Engineer|Backend Engineer|Frontend Engineer|Full Stack Developer|Data Scientist|Business Analyst|Product Manager)\b",
]


def _dedupe(items: List[str]) -> List[str]:
    seen = set()
    output = []
    for item in items:
        cleaned = item.strip(" -•\t\r\n,.;:")
        if not cleaned:
            continue
        key = cleaned.lower()
        if key not in seen:
            seen.add(key)
            output.append(cleaned)
    return output


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document

        document = Document(path)
        chunks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    chunks.append(cell.text)
        return "\n".join(chunk for chunk in chunks if chunk)
    except Exception:
        return ""


def _extract_section_lines(text: str, anchors: List[str]) -> List[str]:
    lines = [line.strip() for line in text.splitlines()]
    captured = []
    active = False
    section_re = re.compile(r"^[A-Za-z ]{3,35}:?$")

    for line in lines:
        lowered = line.lower().strip(":")
        if any(anchor in lowered for anchor in anchors):
            active = True
            tail = re.sub(r"^[^:]{0,40}:", "", line).strip()
            if tail and tail != line:
                captured.append(tail)
            continue
        if active and section_re.match(line) and not line.startswith(("-", "•")):
            active = False
        elif active and line:
            captured.append(line)

    return captured


def _split_skill_lines(lines: List[str]) -> List[str]:
    skills = []
    for line in lines:
        parts = re.split(r"[,;/|]|\band\b", line)
        for part in parts:
            cleaned = re.sub(r"^[\-•*]\s*", "", part).strip()
            if (
                1 <= len(cleaned.split()) <= 4
                and len(cleaned) <= 32
                and not re.search(r"\b(this|that|beyond|practical|probably|terms|range|used)\b", cleaned, re.IGNORECASE)
            ):
                skills.append(cleaned)
    return _dedupe(skills)


def _known_skill_hits(text: str) -> List[str]:
    hits = []
    lowered = text.lower()
    for skill in KNOWN_SKILLS:
        if skill.lower() in lowered:
            hits.append(skill)
    return _dedupe(hits)


def _extract_title(text: str) -> str:
    compact = " ".join(text.split())
    for pattern in TITLE_PATTERNS:
        match = re.search(pattern, compact, flags=re.IGNORECASE)
        if match:
            return match.group(1).strip()
    return DEFAULT_JD["target_title"]


def _extract_experience(text: str) -> int:
    match = re.search(r"(\d+)\+?\s*(?:years?|yrs?)", text, flags=re.IGNORECASE)
    if not match:
        return 0
    return int(match.group(1))


def _extract_industry(text: str) -> str:
    match = re.search(r"(?:industry|domain)[:\s]+([A-Za-z &/-]{3,40})", text, flags=re.IGNORECASE)
    if not match:
        return DEFAULT_JD["target_industry"]
    return match.group(1).strip(" .")


def _extract_field(text: str) -> str:
    fields = [
        "Computer Science",
        "Data Science",
        "Information Technology",
        "Statistics",
        "Mathematics",
        "Engineering",
        "Business",
        "Design",
        "Marketing",
    ]
    lowered = text.lower()
    for field in fields:
        if field.lower() in lowered:
            return field
    return DEFAULT_JD["target_field"]


def parse_jd_text(text: str) -> Dict[str, object]:
    """Parse JD text into the stable dict expected by the scorer."""
    text = text or ""
    required = _split_skill_lines(_extract_section_lines(text, ["required", "must have", "key qualification"]))
    preferred = _split_skill_lines(_extract_section_lines(text, ["preferred", "nice to have", "good to have"]))

    known_hits = _known_skill_hits(text)
    required = _dedupe([skill for skill in required if skill.lower() in {hit.lower() for hit in known_hits}])
    preferred = _dedupe([skill for skill in preferred if skill.lower() in {hit.lower() for hit in known_hits}])

    if not required:
        required = known_hits[:8]
    else:
        preferred = _dedupe(preferred + [skill for skill in known_hits if skill.lower() not in {item.lower() for item in required}])

    target_title = _extract_title(text)
    target_industry = _extract_industry(text)
    skills_text = " ".join(_dedupe(required + preferred)) + f" {target_title} {target_industry}"

    return {
        "required_skills": required,
        "preferred_skills": preferred,
        "target_title": target_title,
        "min_experience_years": _extract_experience(text),
        "target_industry": target_industry,
        "target_field": _extract_field(text),
        "skills_text": skills_text.strip() or text,
    }


def parse_jd_docx(path: str) -> Dict[str, object]:
    """Parse a .docx JD. Safe defaults are returned if parsing fails."""
    jd_path = Path(path)
    text = _extract_docx_text(jd_path) if jd_path.exists() else ""
    parsed = parse_jd_text(text)
    for key, value in DEFAULT_JD.items():
        parsed.setdefault(key, value)
    return parsed
